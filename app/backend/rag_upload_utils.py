import logging
from typing import List
import io
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

def extract_text(filename: str, raw: bytes) -> str:
    """
    Extracts plain text from a file based on its filename extension.
    Requires additional libraries like PyPDF2, python-docx.
    """
    logger.info(f"Extracting text from {filename}...")
    # Placeholder: Implement actual extraction based on file type
    # Example for plain text:
    if filename.lower().endswith(".txt"):
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            logger.warning(f"Could not decode {filename} as UTF-8, trying latin-1")
            return raw.decode('latin-1') # Fallback
    elif filename.lower().endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text)
            return "\n".join(pages)
        except Exception as e:
            logger.error(f"Error extracting PDF text from {filename}: {e}")
            return ""
    elif filename.lower().endswith((".doc", ".docx")):
        # try:
            # 更加健壮的DOCX处理
            doc = Document(io.BytesIO(raw))
            full_text = []
            
            # 提取所有段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            # 确保有内容被提取
            if not full_text:
                logger.warning(f"No content extracted from {filename}")
                return ""
                
            result = "\n".join(full_text)
            logger.info(f"Successfully extracted {len(result)} characters from {filename}")
            return result
        # except Exception as e:
        #     logger.error(f"Error extracting DOCX text from {filename}: {e}", exc_info=True)
        #     return ""
    else:
        logger.warning(f"Unsupported file type for extraction: {filename}")
        return ""

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    Splits text into chunks of a specified size with overlap.
    """
    logger.info(f"Chunking text (length {len(text)})...")
    if not text:
        return []
    # Placeholder: Implement a simple chunking strategy
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap # Move start pointer with overlap
        if start >= len(text): # Avoid infinite loop if overlap >= chunk_size
            break
    logger.info(f"Created {len(chunks)} chunks.")
    return chunks
